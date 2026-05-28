"""Document text extraction helpers."""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.rag.document_validation import CONTENT_TYPE_DOCX, CONTENT_TYPE_MARKDOWN, CONTENT_TYPE_PDF


SUPPORTED_TEXT_TYPES = {
    "text/plain",
    CONTENT_TYPE_MARKDOWN,
    "application/x-markdown",
}

SUPPORTED_EXTRACTION_TYPES = {
    *SUPPORTED_TEXT_TYPES,
    CONTENT_TYPE_PDF,
    CONTENT_TYPE_DOCX,
}


class OcrRequiredError(ValueError):
    """Raised when a scanned document needs OCR that is not enabled."""


def extract_text(
    content: bytes,
    content_type: str | None = None,
    *,
    max_pages: int | None = None,
    ocr_enabled: bool = False,
) -> str:
    """Extract text from supported document bytes."""
    normalized_type = (content_type or "text/plain").split(";")[0].strip().lower()
    if normalized_type == "application/x-markdown":
        normalized_type = CONTENT_TYPE_MARKDOWN
    if normalized_type in SUPPORTED_TEXT_TYPES:
        return content.decode("utf-8")
    if normalized_type == CONTENT_TYPE_PDF:
        return extract_pdf_text(content, max_pages=max_pages, ocr_enabled=ocr_enabled)
    if normalized_type == CONTENT_TYPE_DOCX:
        return extract_docx_text(content)
    raise ValueError(f"Unsupported content type for document ingestion: {content_type}")


def extract_pdf_text(content: bytes, *, max_pages: int | None, ocr_enabled: bool) -> str:
    """Extract selectable text from a PDF and gate scanned/OCR documents."""
    reader = PdfReader(BytesIO(content))
    page_count = len(reader.pages)
    if max_pages is not None and page_count > max_pages:
        raise ValueError("PDF exceeds the configured page limit")
    page_text = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(part for part in page_text if part)
    if text.strip():
        return text
    if not ocr_enabled:
        raise OcrRequiredError("PDF appears scanned; OCR is required but not enabled")
    raise OcrRequiredError("OCR processing path is configured but not implemented in this build")


def extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX paragraphs and tables."""
    document = DocxDocument(BytesIO(content))
    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)

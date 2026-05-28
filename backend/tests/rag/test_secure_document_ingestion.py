from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

import app.models  # noqa: F401
from app.core.config import Settings, get_settings
from app.core.passwords import hash_password
from app.core.rate_limit import rate_limiter
from app.db.base import Base, utc_now
from app.db.session import get_db_session
from app.jobs.processor import BackgroundJobProcessor
from app.jobs.service import (
    JOB_STATUS_COMPLETED,
    JOB_TYPE_RAG_DOCUMENT_FILE_INGESTION,
    BackgroundJobService,
)
from app.main import create_app
from app.models import BackgroundJob, BusinessUser, DocumentChunk, KnowledgeDocument
from app.rag.document_storage import LocalDocumentStorage
from app.rag.document_validation import (
    DocumentValidationError,
    validate_document_upload,
)
from app.rag.extraction import OcrRequiredError, extract_text
from app.rag.ingestion import RagIngestionService
from app.tenants.service import TenantService
from fastapi.testclient import TestClient


class NoopWakeQueue:
    def notify(self, _job_id: str) -> bool:
        return True


def create_test_session_factory():
    engine = create_engine(
        "sqlite:///:memory:",
        future=True,
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    return sessionmaker(bind=engine, future=True, expire_on_commit=False)


def seed_business_user(session):
    tenant_service = TenantService(session)
    tenant = tenant_service.create_tenant("Demo Plumbing", "demo-plumbing")
    business = tenant_service.create_business(
        tenant_id=tenant.id,
        name="Demo Plumbing",
        email="owner@example.test",
    )
    user = BusinessUser(
        tenant_id=tenant.id,
        business_id=business.id,
        email="owner@example.test",
        full_name="Demo Owner",
        role="owner",
        status="active",
        password_hash=hash_password("correct-password"),
        password_updated_at=utc_now(),
    )
    session.add(user)
    session.flush()
    return tenant, user


def create_client(session, monkeypatch, tmp_path: Path):
    monkeypatch.setenv("AI_PROVIDER", "test")
    monkeypatch.setenv("AI_EMBEDDING_DIMENSIONS", "16")
    monkeypatch.setenv("BUSINESS_PORTAL_SESSION_SECRET", "test-secret")
    monkeypatch.setenv("DOCUMENT_STORAGE_ROOT", str(tmp_path / "documents"))
    monkeypatch.setattr("app.jobs.redis_queue.RedisWakeQueue.notify", lambda *_args: True)
    rate_limiter.reset()
    get_settings.cache_clear()
    app = create_app()

    def override_db():
        yield session

    app.dependency_overrides[get_db_session] = override_db
    return TestClient(app)


def login(client: TestClient) -> str:
    response = client.post(
        "/business-portal/auth/login",
        json={
            "tenant_slug": "demo-plumbing",
            "email": "owner@example.test",
            "password": "correct-password",
        },
    )
    assert response.status_code == 200
    return response.json()["access_token"]


def auth_header(token: str) -> dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


def build_docx_bytes(text: str) -> bytes:
    from io import BytesIO

    buffer = BytesIO()
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


def build_blank_pdf_bytes() -> bytes:
    from io import BytesIO

    buffer = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(buffer)
    return buffer.getvalue()


def test_document_upload_validation_rejects_unsafe_inputs(tmp_path):
    settings = Settings(document_upload_max_bytes=8, document_storage_root=str(tmp_path))

    with pytest.raises(DocumentValidationError, match="Unsupported document extension"):
        validate_document_upload(
            filename="payload.exe",
            content=b"123",
            content_type="application/octet-stream",
            settings=settings,
        )
    with pytest.raises(DocumentValidationError, match="size limit"):
        validate_document_upload(
            filename="services.txt",
            content=b"123456789",
            content_type="text/plain",
            settings=settings,
        )
    with pytest.raises(DocumentValidationError, match="signature"):
        validate_document_upload(
            filename="brochure.pdf",
            content=b"not pdf",
            content_type="application/pdf",
            settings=settings,
        )
    with pytest.raises(DocumentValidationError, match="malformed"):
        validate_document_upload(
            filename="brochure.pdf",
            content=b"%PDF-1.4 broken",
            content_type="application/pdf",
            settings=Settings(document_upload_max_bytes=1024, document_storage_root=str(tmp_path)),
        )
    with pytest.raises(DocumentValidationError, match="malware"):
        validate_document_upload(
            filename="services.txt",
            content=b"X5O!P%@AP[4\\PZX54(P^)7CC)7}$EICAR",
            content_type="text/plain",
            settings=Settings(document_upload_max_bytes=1024, document_storage_root=str(tmp_path)),
        )


def test_document_validation_sanitizes_filename_and_accepts_docx(tmp_path):
    content = build_docx_bytes("Blocked drains and hot water repairs.")
    settings = Settings(document_storage_root=str(tmp_path))

    validated = validate_document_upload(
        filename="../Demo Services.docx",
        content=content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        settings=settings,
    )
    extracted = extract_text(content, validated.content_type)

    assert validated.filename == "Demo Services.docx"
    assert validated.malware_scan_status == "passed_basic"
    assert "Blocked drains" in extracted

    markdown = validate_document_upload(
        filename="services.md",
        content=b"# Services\nBlocked drains",
        content_type="text/plain",
        settings=settings,
    )
    assert markdown.content_type == "text/markdown"


def test_pdf_without_selectable_text_is_gated_for_ocr():
    with pytest.raises(OcrRequiredError, match="OCR is required"):
        extract_text(
            build_blank_pdf_bytes(),
            "application/pdf",
            max_pages=10,
            ocr_enabled=False,
        )


def test_stored_file_job_processes_without_sensitive_payload(tmp_path):
    session_factory = create_test_session_factory()
    settings = Settings(
        ai_provider="test",
        ai_embedding_dimensions=16,
        rag_chunk_size=6,
        rag_chunk_overlap=2,
        document_storage_root=str(tmp_path / "documents"),
    )
    with session_factory() as session:
        tenant = TenantService(session).create_tenant("Demo Plumbing", "demo-plumbing")
        content = b"Emergency blocked drains and hot water repairs in Bondi"
        validated = validate_document_upload(
            filename="services.txt",
            content=content,
            content_type="text/plain",
            settings=settings,
        )
        document = RagIngestionService(
            session=session,
            embedding_provider=None,
            settings=settings,
        ).create_pending_document(
            tenant_id=tenant.id,
            filename=validated.filename,
            content_type=validated.content_type,
            source_type="uploaded_file",
            file_size_bytes=validated.size_bytes,
            file_sha256=validated.sha256,
            malware_scan_status=validated.malware_scan_status,
        )
        document.storage_path = LocalDocumentStorage(settings).save(
            tenant_id=tenant.id,
            document_id=document.id,
            filename=validated.filename,
            content=content,
        )
        job = BackgroundJobService(
            session,
            settings,
            redis_wake_queue=NoopWakeQueue(),
        ).enqueue_document_file_ingestion(tenant_id=tenant.id, document_id=document.id).job
        session.commit()
        job_id = job.id

    processed_job = BackgroundJobProcessor(
        session_factory,
        settings=settings,
        worker_id="test-worker",
    ).process_one()

    with session_factory() as session:
        stored_job = session.get(BackgroundJob, job_id)
        document = session.scalars(select(KnowledgeDocument)).one()
        chunks = list(session.scalars(select(DocumentChunk)))

        assert processed_job is not None
        assert stored_job.status == JOB_STATUS_COMPLETED
        assert stored_job.job_type == JOB_TYPE_RAG_DOCUMENT_FILE_INGESTION
        assert "content" not in stored_job.payload
        assert document.status == "ingested"
        assert document.extraction_status == "extracted"
        assert chunks


def test_business_portal_file_upload_refresh_delete_are_tenant_scoped(monkeypatch, tmp_path):
    session_factory = create_test_session_factory()
    with session_factory() as session:
        tenant, _user = seed_business_user(session)
        other_tenant = TenantService(session).create_tenant("Other Trade", "other-trade")
        other_doc = KnowledgeDocument(
            tenant_id=other_tenant.id,
            filename="other.txt",
            status="ingested",
        )
        session.add(other_doc)
        session.commit()
        client = create_client(session, monkeypatch, tmp_path)
        token = login(client)

        upload_response = client.post(
            "/business-portal/documents/upload",
            headers=auth_header(token),
            files={"file": ("services.txt", b"Blocked drains Bondi", "text/plain")},
        )
        document_id = upload_response.json()["id"]
        refresh_response = client.post(
            f"/business-portal/documents/{document_id}/refresh",
            headers=auth_header(token),
        )
        cross_tenant_refresh = client.post(
            f"/business-portal/documents/{other_doc.id}/refresh",
            headers=auth_header(token),
        )
        delete_response = client.delete(
            f"/business-portal/documents/{document_id}",
            headers=auth_header(token),
        )

        assert upload_response.status_code == 200
        assert upload_response.json()["source_type"] == "uploaded_file"
        assert upload_response.json()["malware_scan_status"] == "passed_basic"
        assert upload_response.json()["file_size_bytes"] == len(b"Blocked drains Bondi")
        assert upload_response.json()["job_id"]
        assert refresh_response.status_code == 200
        assert refresh_response.json()["status"] == "queued"
        assert cross_tenant_refresh.status_code == 404
        assert delete_response.status_code == 204
        assert session.get(KnowledgeDocument, document_id) is None
        assert session.get(KnowledgeDocument, other_doc.id) is not None
        assert tenant.id != other_tenant.id
